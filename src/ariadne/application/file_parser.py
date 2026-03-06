"""File parser service for extracting text from various file formats."""

from __future__ import annotations

import io
import logging
import sys
import warnings
from pathlib import Path
from typing import Optional

from ariadne.infrastructure.app_logger import get_logger

# Suppress ALL pdfplumber and pdfminer warnings
logging.getLogger("pdfplumber").setLevel(logging.CRITICAL)
logging.getLogger("pdfminer").setLevel(logging.CRITICAL)
logging.getLogger("pdfminer.psparser").setLevel(logging.CRITICAL)

# Suppress warnings from pdfminer about fonts
warnings.filterwarnings("ignore", message=".*FontBBox.*")
warnings.filterwarnings("ignore", message=".*pdfminer.*")

logger = get_logger("file_parser")


def _flush_logs():
    """Force flush all log handlers."""
    for handler in logging.getLogger().handlers:
        handler.flush()
    sys.stdout.flush()
    sys.stderr.flush()


class FileParserService:
    """Service for parsing different file formats and extracting text content."""

    # Maximum file size to read into memory (50MB)
    MAX_READ_SIZE = 50 * 1024 * 1024

    def parse_pdf(self, file_path: str) -> str:
        """
        Parse PDF file and extract text content.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            ImportError: If pdfplumber is not installed
        """
        try:
            import pdfplumber
        except ImportError as exc:
            raise ImportError(
                "pdfplumber is required for PDF parsing. "
                "Install it with: pip install pdfplumber"
            ) from exc

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        text_parts = []
        try:
            with pdfplumber.open(path) as pdf:
                total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            text_parts.append(page_text)
                        logger.debug(
                            "Parsed PDF page %d/%d, chars=%d",
                            page_num,
                            total_pages,
                            len(page_text),
                        )
                    except Exception as exc:  # noqa: BLE001
                        logger.warning("Failed to parse PDF page %d: %s", page_num, exc)

            result = "\n\n".join(text_parts)
            logger.info(
                "PDF parsed: pages=%d, total_chars=%d",
                total_pages,
                len(result),
            )
            return result
        except Exception as exc:  # noqa: BLE001
            logger.error("Failed to parse PDF file %s: %s", file_path, exc)
            raise

    def parse_docx(self, file_path: str) -> str:
        """
        Parse DOCX file and extract text content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            ImportError: If python-docx is not installed
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            ) from exc

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            doc = Document(path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            result = "\n\n".join(text_parts)
            logger.info(
                "DOCX parsed: paragraphs=%d, tables=%d, total_chars=%d",
                len(doc.paragraphs),
                len(doc.tables),
                len(result),
            )
            return result
        except Exception as exc:  # noqa: BLE001
            logger.error("Failed to parse DOCX file %s: %s", file_path, exc)
            raise

    def parse_txt(self, file_path: str) -> str:
        """
        Parse plain text file.

        Args:
            file_path: Path to the text file

        Returns:
            File content as string

        Raises:
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            # Try UTF-8 first, fallback to common encodings
            encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"]
            content = None
            for encoding in encodings:
                try:
                    content = path.read_text(encoding=encoding)
                    break
                except (UnicodeDecodeError, LookupError):
                    continue

            if content is None:
                content = path.read_text(encoding="utf-8", errors="replace")

            logger.info(
                "TXT parsed: encoding=%s, total_chars=%d",
                "detected" if content else "utf-8-fallback",
                len(content),
            )
            return content
        except Exception as exc:  # noqa: BLE001
            logger.error("Failed to parse TXT file %s: %s", file_path, exc)
            raise

    def parse_markdown(self, file_path: str) -> str:
        """
        Parse Markdown file.

        Args:
            file_path: Path to the Markdown file

        Returns:
            File content as string (Markdown formatting preserved)
        """
        # Markdown is essentially text with formatting, so we can use TXT parsing
        return self.parse_txt(file_path)

    def parse(self, file_path: str, file_type: str) -> str:
        """
        Parse file based on its type.

        Args:
            file_path: Path to the file
            file_type: File extension (pdf, docx, txt, md)

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported
        """
        file_type = file_type.lower().lstrip(".")

        parsers = {
            "pdf": self.parse_pdf,
            "docx": self.parse_docx,
            "txt": self.parse_txt,
            "md": self.parse_markdown,
            "markdown": self.parse_markdown,
        }

        parser_func = parsers.get(file_type)
        if not parser_func:
            raise ValueError(
                f"Unsupported file type: {file_type}. "
                f"Supported types: {', '.join(parsers.keys())}"
            )

        logger.info("Parsing file: type=%s, path=%s", file_type, file_path)
        return parser_func(file_path)

    def parse_from_bytes(self, content: bytes, file_type: str, file_name: str) -> str:
        """
        Parse file from bytes content.

        This is useful for handling uploaded files without saving to disk first.

        Args:
            content: File content as bytes
            file_type: File extension (pdf, docx, txt, md)
            file_name: Original file name (for logging)

        Returns:
            Extracted text content
        """
        import tempfile

        file_type = file_type.lower().lstrip(".")

        # For text files, decode directly
        if file_type in {"txt", "md", "markdown"}:
            encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"]
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except (UnicodeDecodeError, LookupError):
                    continue
            return content.decode("utf-8", errors="replace")

        # For binary files (PDF, DOCX), create temporary file
        with tempfile.NamedTemporaryFile(
            suffix=f".{file_type}",
            delete=True,
        ) as tmp_file:
            tmp_file.write(content)
            tmp_file.flush()
            return self.parse(tmp_file.name, file_type)

    def get_preview(self, content: str, max_length: int = 200) -> str:
        """
        Get a preview of the text content.

        Args:
            content: Full text content
            max_length: Maximum length of preview

        Returns:
            Preview text (truncated if necessary)
        """
        if not content:
            return ""
        if len(content) <= max_length:
            return content
        return content[:max_length] + "..."
